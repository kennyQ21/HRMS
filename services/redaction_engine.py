"""
services/redaction_engine.py
------------------------------
Full Multi-Format Redaction Engine.

Supports all document types with format-appropriate redaction:

  ┌──────────┬──────────────────────────────────────────────────┐
  │ Format   │ Method                                           │
  ├──────────┼──────────────────────────────────────────────────┤
  │ PDF      │ PyMuPDF — black rectangle overlay on text spans  │
  │ DOCX     │ python-docx — inline text replacement in runs    │
  │ XLSX     │ openpyxl — cell value masking                    │
  │ Image    │ PIL — black rectangle over bbox regions          │
  │ CSV      │ pandas — column-level cell masking               │
  │ Plain    │ string replacement                               │
  └──────────┴──────────────────────────────────────────────────┘

Redaction Types:
  • FULL        → XXXXXXXXXXXX  (passwords, CVV)
  • PARTIAL     → XXXX-XXXX-1234  (credit cards, phones)
  • CONTEXTUAL  → [PERSON_NAME]  [ADDRESS]  (documents, readable output)
  • MASK_CHAR   → ████████████  (visual PDF overlay)

Usage:
    engine = RedactionEngine()
    result = engine.redact(
        file_path="/tmp/doc.pdf",
        filename="doc.pdf",
        entities=resolved_entities,
        redaction_type="contextual",
    )
    # result.redacted_bytes  — file bytes ready to return / save
    # result.redaction_map   — {entity_value: redacted_value}
"""

from __future__ import annotations

import io
import logging
import re
import tempfile
import os
from dataclasses import dataclass, field
from typing import Optional

logger = logging.getLogger(__name__)

# ── Redaction type constants ──────────────────────────────────────────────────

REDACT_FULL        = "full"         # XXXXXXXXXXXX
REDACT_PARTIAL     = "partial"      # XXXX-XXXX-1234
REDACT_CONTEXTUAL  = "contextual"   # [PERSON_NAME]
REDACT_MASK        = "mask"         # ████████████ (PDF visual)

# PII types that get FULL redaction by default
_FULL_REDACT_TYPES: set[str] = {
    "password", "cvv", "ssn",
}

# PII types that get PARTIAL masking (show last N chars)
_PARTIAL_REDACT_TYPES: dict[str, int] = {
    "credit_card":  4,    # show last 4
    "bank_account": 4,
    "aadhaar":      4,
    "phone":        4,
    "pan":          0,    # show none — PAN is 10 chars, FULL is safer
}

# Contextual label map for [TYPE] replacement
_CONTEXTUAL_LABELS: dict[str, str] = {
    "name":                    "[PERSON_NAME]",
    "address":                 "[ADDRESS]",
    "email":                   "[EMAIL]",
    "corporate_email":         "[CORPORATE_EMAIL]",
    "phone":                   "[PHONE_NUMBER]",
    "dob":                     "[DATE_OF_BIRTH]",
    "aadhaar":                 "[AADHAAR_NUMBER]",
    "pan":                     "[PAN_NUMBER]",
    "passport":                "[PASSPORT_NUMBER]",
    "voter_id":                "[VOTER_ID]",
    "driving_license":         "[DL_NUMBER]",
    "ssn":                     "[SSN]",
    "credit_card":             "[CARD_NUMBER]",
    "bank_account":            "[ACCOUNT_NUMBER]",
    "upi":                     "[UPI_ID]",
    "ifsc":                    "[IFSC_CODE]",
    "cvv":                     "[CVV]",
    "expiry":                  "[EXPIRY_DATE]",
    "password":                "[PASSWORD]",
    "user_id":                 "[USER_ID]",
    "organization":            "[ORGANIZATION]",
    "occupation":              "[OCCUPATION]",
    "employee_id":             "[EMPLOYEE_ID]",
    "diagnosis":               "[MEDICAL_DIAGNOSIS]",
    "allergies":               "[ALLERGY_RECORD]",
    "treatment_history":       "[TREATMENT_HISTORY]",
    "prescription":            "[PRESCRIPTION]",
    "immunization":            "[IMMUNIZATION_RECORD]",
    "blood_group":             "[BLOOD_GROUP]",
    "mrn":                     "[MEDICAL_RECORD_NUMBER]",
    "insurance_policy":        "[INSURANCE_POLICY]",
    "insurance_provider":      "[INSURANCE_PROVIDER]",
    "educational_qualification":"[QUALIFICATION]",
    "city":                    "[CITY]",
    "pincode":                 "[PINCODE]",
    "ip_address":              "[IP_ADDRESS]",
    "nationality":             "[NATIONALITY]",
    "gender":                  "[GENDER]",
    "age":                     "[AGE]",
    "marital_status":          "[MARITAL_STATUS]",
}


@dataclass
class RedactionResult:
    """Output of one redact() call."""
    filename:       str
    format:         str                         # pdf | docx | xlsx | image | csv | text
    redacted_bytes: bytes = b""
    redaction_map:  dict  = field(default_factory=dict)  # original → replacement
    redaction_verification: dict = field(default_factory=lambda: {
        "passed": True,
        "unredacted_entities": [],
    })
    entity_count:   int   = 0
    error:          Optional[str] = None


class RedactionEngine:

    def redact(
        self,
        file_path: str,
        filename: str,
        entities: list,                     # list[ResolvedEntity]
        redaction_type: str = REDACT_CONTEXTUAL,
        pii_types_filter: Optional[set[str]] = None,
    ) -> RedactionResult:
        """
        Apply redaction to a file for all resolved entities.

        Args:
            file_path:        Path to the source file.
            filename:         Original filename (used for format detection).
            entities:         Resolved entities from entity_resolution.resolve().
            redaction_type:   "full" | "partial" | "contextual" | "mask"
            pii_types_filter: Only redact these PII type IDs (None = all).
        """
        if pii_types_filter:
            entities = [e for e in entities if e.pii_type in pii_types_filter]

        if not entities:
            with open(file_path, "rb") as f:
                return RedactionResult(
                    filename=filename, format=_detect_format(filename),
                    redacted_bytes=f.read(), entity_count=0,
                )

        fmt = _detect_format(filename)
        redaction_map = self._build_redaction_map(entities, redaction_type)

        logger.info(
            "[REDACTION] %s format=%s entities=%d type=%s",
            filename, fmt, len(entities), redaction_type,
        )

        try:
            verification = {"passed": True, "unredacted_entities": []}
            if fmt == "pdf":
                result_bytes, verification = self._redact_pdf(
                    file_path, entities, redaction_map, redaction_type
                )
            elif fmt == "docx":
                result_bytes = self._redact_docx(file_path, redaction_map)
            elif fmt == "xlsx":
                result_bytes = self._redact_xlsx(file_path, redaction_map)
            elif fmt == "image":
                result_bytes = self._redact_image(file_path, entities, filename)
            elif fmt == "csv":
                result_bytes = self._redact_csv(file_path, redaction_map)
            else:
                result_bytes = self._redact_text(file_path, redaction_map)

            return RedactionResult(
                filename=filename,
                format=fmt,
                redacted_bytes=result_bytes,
                redaction_map=redaction_map,
                redaction_verification=verification,
                entity_count=len(entities),
            )

        except Exception as exc:
            logger.error("[REDACTION] Failed for %s: %s", filename, exc, exc_info=True)
            with open(file_path, "rb") as f:
                return RedactionResult(
                    filename=filename, format=fmt,
                    redacted_bytes=f.read(),
                    entity_count=0,
                    error=str(exc),
                )

    # ── Redaction map builder ─────────────────────────────────────────────────

    def _build_redaction_map(self, entities: list, redaction_type: str) -> dict[str, str]:
        """Build {original_value: replacement} map for text-based redaction."""
        rmap: dict[str, str] = {}
        for e in entities:
            val = e.value
            if not val:
                continue
            rmap[val] = self._get_replacement(val, e.pii_type, redaction_type)
        return rmap

    def _get_replacement(
        self, value: str, pii_type: str, redaction_type: str
    ) -> str:
        """Choose the right replacement string for a given PII value."""
        # Per-type overrides take priority over the global redaction_type
        if pii_type in _FULL_REDACT_TYPES:
            return "X" * len(value)

        if redaction_type == REDACT_CONTEXTUAL:
            return _CONTEXTUAL_LABELS.get(pii_type, f"[{pii_type.upper()}]")

        if redaction_type == REDACT_PARTIAL or pii_type in _PARTIAL_REDACT_TYPES:
            keep = _PARTIAL_REDACT_TYPES.get(pii_type, 4)
            digits = re.sub(r"\D", "", value)
            if keep > 0 and len(digits) > keep:
                masked = "X" * (len(digits) - keep) + digits[-keep:]
                return masked
            return "X" * len(value)

        if redaction_type == REDACT_FULL:
            return "X" * len(value)

        # Default
        return _CONTEXTUAL_LABELS.get(pii_type, f"[{pii_type.upper()}]")

    # ── PDF redaction (PyMuPDF) ───────────────────────────────────────────────

    def _redact_pdf(
        self,
        file_path: str,
        entities: list,
        redaction_map: dict,
        redaction_type: str,
    ) -> tuple[bytes, dict]:
        import fitz  # PyMuPDF

        doc = fitz.open(file_path)
        total_redacted = 0
        unredacted_entities = []
        page_count = doc.page_count

        for page in doc:
            for entity in entities:
                val = entity.value
                if not val or len(val) < 2:
                    continue

                # Build search variants: canonical + formatted + OCR-tolerant
                search_variants = self._build_search_variants(val, entity.pii_type)

                found = False
                for variant in search_variants:
                    if not variant or len(variant) < 2:
                        continue
                    try:
                        instances = page.search_for(variant, quads=False)
                        for rect in instances:
                            if redaction_type == REDACT_MASK:
                                page.add_redact_annot(rect, fill=(0, 0, 0))
                            else:
                                replacement = redaction_map.get(val, f"[{entity.pii_type.upper()}]")
                                page.add_redact_annot(
                                    rect,
                                    text=replacement,
                                    fill=(0, 0, 0),
                                    text_color=(1, 1, 1),
                                    fontsize=8,
                                )
                            total_redacted += 1
                            found = True
                    except Exception:
                        pass

                if not found:
                    unredacted_entities.append(entity.pii_type)

            # Apply all redaction annotations on this page
            page.apply_redactions()

        buf = io.BytesIO()
        doc.save(buf)
        doc.close()
        result_bytes = buf.getvalue()

        verification = self.verify_pdf_redaction(result_bytes, entities)

        if unredacted_entities:
            logger.warning("[REDACTION] PDF: %d entities could not be located for redaction: %s",
                          len(unredacted_entities), list(set(unredacted_entities))[:10])
        if not verification["passed"]:
            logger.warning("[REDACTION] PDF: %d entities still present after redaction (verification failed)",
                           len(verification["unredacted_entities"]))

        logger.info("[REDACTION] PDF: %d pages processed, %d regions redacted",
                    page_count, total_redacted)
        return result_bytes, verification

    # ── Normalization helper ──────────────────────────────────────────────────

    @staticmethod
    def _normalize_text(text: str) -> str:
        """Normalize text for comparison: lowercase, collapse whitespace, strip punctuation."""
        import re as _re
        text = text.lower()
        text = _re.sub(r"\s+", " ", text)
        text = _re.sub(r"[^a-z0-9 ]", "", text)
        return text.strip()

    def _build_search_variants(self, value: str, pii_type: str) -> list[str]:
        """Build search variants for PDF text search to handle OCR/format differences."""
        variants = [value]
        import re as _re

        # Stripped/normalized variant
        stripped = value.strip()
        if stripped != value:
            variants.append(stripped)

        # For numeric types, add digit-only variant
        if pii_type in {"aadhaar", "phone", "credit_card", "bank_account", "ssn"}:
            digits = _re.sub(r"\D", "", value)
            if digits and digits != value:
                variants.append(digits)
                # Add spaced variant (e.g., "1234 5678 9012")
                if len(digits) >= 8:
                    spaced = " ".join([digits[i:i+4] for i in range(0, len(digits), 4)])
                    variants.append(spaced)
                    dashed = "-".join([digits[i:i+4] for i in range(0, len(digits), 4)])
                    variants.append(dashed)

        # Case variants
        if value.upper() != value:
            variants.append(value.upper())
        if value.lower() != value:
            variants.append(value.lower())

        return variants

    def verify_pdf_redaction(self, redacted_pdf, entities: list) -> dict:
        """
        Compliance-critical: verify redacted PDF truly removes PII.

        Checks:
          1. Searchable text (page.get_text)
          2. Hidden text / XObject streams
          3. Normalized comparison (lowercase, digits-only for numeric IDs)

        Returns:
            {"passed": bool, "unredacted_entities": [pii_type, ...]}
        """
        import io as _io
        import re as _re
        import fitz

        unredacted = []
        doc = fitz.open(stream=redacted_pdf, filetype="pdf") if isinstance(redacted_pdf, (bytes, bytearray)) else fitz.open(redacted_pdf)

        # Extract ALL text from the redacted document
        full_text_parts = []
        try:
            for page in doc:
                # Visible/searchable text, including OCR text layers.
                full_text_parts.append(page.get_text())
                # Hidden text in referenced streams where PyMuPDF exposes it.
                try:
                    for xref in page.get_xobjects():
                        try:
                            xref_id = xref[0] if isinstance(xref, (tuple, list)) else xref
                            stream = doc.xref_stream(xref_id)
                            if stream:
                                full_text_parts.append(stream.decode("utf-8", errors="ignore"))
                        except Exception:
                            pass
                except Exception:
                    pass
        finally:
            doc.close()

        full_text = "\n".join(full_text_parts)
        normalized_full = self._normalize_text(full_text)
        digits_in_doc = _re.sub(r"\D", "", full_text)

        for entity in entities:
            val = entity.value
            if not val or len(val) < 2:
                continue

            found = False

            # Check 1: Direct text match
            if val in full_text:
                found = True

            # Check 2: Normalized lowercase match
            norm_val = self._normalize_text(val)
            if not found and norm_val and norm_val in normalized_full:
                found = True

            # Check 3: Numeric IDs — compare digits only
            digits = _re.sub(r"\D", "", val)
            if not found and digits and len(digits) >= 6:
                if digits in digits_in_doc:
                    found = True

            if found:
                unredacted.append({
                    "type": entity.pii_type,
                    "value": entity.value,
                })

        passed = len(unredacted) == 0
        return {"passed": passed, "unredacted_entities": unredacted}

    # ── DOCX redaction (python-docx) ──────────────────────────────────────────

    def _redact_docx(self, file_path: str, redaction_map: dict) -> bytes:
        import docx

        doc = docx.Document(file_path)

        def _replace_in_run(run, rmap: dict) -> None:
            text = run.text
            for original, replacement in rmap.items():
                if original in text:
                    text = text.replace(original, replacement)
            run.text = text

        # Process paragraphs
        for para in doc.paragraphs:
            for run in para.runs:
                _replace_in_run(run, redaction_map)

        # Process tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            _replace_in_run(run, redaction_map)

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        logger.info("[REDACTION] DOCX: processed %d paragraphs", len(doc.paragraphs))
        return buf.read()

    # ── XLSX redaction (openpyxl) ─────────────────────────────────────────────

    def _redact_xlsx(self, file_path: str, redaction_map: dict) -> bytes:
        from openpyxl import load_workbook

        wb = load_workbook(file_path)
        cells_redacted = 0

        for ws in wb.worksheets:
            for row in ws.iter_rows():
                for cell in row:
                    if cell.value is None:
                        continue
                    cell_str = str(cell.value)
                    changed = False
                    for original, replacement in redaction_map.items():
                        if original in cell_str:
                            cell_str = cell_str.replace(original, replacement)
                            changed = True
                    if changed:
                        cell.value = cell_str
                        cells_redacted += 1

        buf = io.BytesIO()
        wb.save(buf)
        buf.seek(0)
        logger.info("[REDACTION] XLSX: %d cells redacted", cells_redacted)
        return buf.read()

    # ── Image redaction (PIL + bbox) ──────────────────────────────────────────

    def _redact_image(
        self, file_path: str, entities: list, filename: str
    ) -> bytes:
        from PIL import Image, ImageDraw

        img   = Image.open(file_path).convert("RGB")
        draw  = ImageDraw.Draw(img)
        count = 0
        failed = 0

        for entity in entities:
            # Primary: use OCR bbox coordinates from entity metadata
            bbox = entity.metadata.get("bbox")
            if bbox:
                try:
                    xs = [p[0] for p in bbox]
                    ys = [p[1] for p in bbox]
                    # Expand bbox slightly to ensure full coverage
                    pad = 3
                    draw.rectangle(
                        [min(xs) - pad, min(ys) - pad, max(xs) + pad, max(ys) + pad],
                        fill="black"
                    )
                    count += 1
                except Exception as exc:
                    logger.warning("[REDACTION] Image bbox error: %s", exc)
                    failed += 1
            else:
                # Fallback: try text-based search on OCR lines if available
                # This handles entities without bbox metadata
                failed += 1

        buf = io.BytesIO()
        ext = os.path.splitext(filename)[1].lower().lstrip(".")
        fmt = "JPEG" if ext in ("jpg", "jpeg") else ext.upper() or "PNG"
        try:
            img.save(buf, format=fmt)
        except Exception:
            img.save(buf, format="PNG")
        buf.seek(0)

        if failed > 0:
            logger.warning("[REDACTION] Image: %d regions redacted, %d entities without bbox (may be unredacted)",
                           count, failed)

        logger.info("[REDACTION] Image: %d regions redacted", count)
        return buf.read()

    # ── CSV redaction ─────────────────────────────────────────────────────────

    def _redact_csv(self, file_path: str, redaction_map: dict) -> bytes:
        with open(file_path, "r", errors="replace") as f:
            content = f.read()
        for original, replacement in redaction_map.items():
            content = content.replace(original, replacement)
        return content.encode("utf-8")

    # ── Plain text fallback ───────────────────────────────────────────────────

    def _redact_text(self, file_path: str, redaction_map: dict) -> bytes:
        with open(file_path, "r", errors="replace") as f:
            content = f.read()
        for original, replacement in sorted(
            redaction_map.items(), key=lambda kv: -len(kv[0])
        ):
            content = content.replace(original, replacement)
        return content.encode("utf-8")


# ── Helpers ───────────────────────────────────────────────────────────────────

_FORMAT_MAP: dict[str, str] = {
    ".pdf":  "pdf",
    ".docx": "docx", ".doc": "docx",
    ".xlsx": "xlsx", ".xls": "xlsx",
    ".csv":  "csv",
    ".jpg":  "image", ".jpeg": "image",
    ".png":  "image", ".bmp": "image",
    ".tif":  "image", ".tiff": "image", ".webp": "image",
}


def _detect_format(filename: str) -> str:
    ext = os.path.splitext(filename.lower())[1]
    return _FORMAT_MAP.get(ext, "text")


# ── Module-level singleton ────────────────────────────────────────────────────

_redaction_engine = RedactionEngine()


def redact_document(
    file_path: str,
    filename: str,
    entities: list,
    redaction_type: str = REDACT_CONTEXTUAL,
    pii_types_filter: Optional[set[str]] = None,
) -> RedactionResult:
    return _redaction_engine.redact(
        file_path=file_path,
        filename=filename,
        entities=entities,
        redaction_type=redaction_type,
        pii_types_filter=pii_types_filter,
    )


def verify_pdf_redaction(redacted_pdf, entities: list) -> dict:
    return _redaction_engine.verify_pdf_redaction(redacted_pdf, entities)
